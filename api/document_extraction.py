"""
Transient document text extraction for résumé and job-description uploads.

Deterministic only: no AI, no OCR, no persistence.
"""

from __future__ import annotations

import json
import re
import sys
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

_SRC_DIR = Path(__file__).resolve().parent.parent / "src"
if str(_SRC_DIR) not in sys.path:
    sys.path.insert(0, str(_SRC_DIR))

from analysis_runner import load_json_file
from extract_keywords import find_skills

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024
MAX_PDF_PAGES = 20
MAX_EXTRACTED_TEXT_LENGTH = 100_000

SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".txt", ".md"})
EXTENSION_TO_SOURCE_KIND = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
    ".md": "md",
}

_DEFAULT_TAXONOMY = Path("data/skills_taxonomy.json")
_DEFAULT_ALIASES = Path("data/skill_aliases.json")

_HEADING_IGNORE = frozenset(
    {
        "resume",
        "résumé",
        "curriculum vitae",
        "cv",
        "experience",
        "education",
        "skills",
        "summary",
        "objective",
        "projects",
        "certifications",
        "contact",
    }
)

_FILENAME_GENERIC_WORDS = re.compile(
    r"\b(resume|résumé|cv|final|updated|copy)\b",
    re.IGNORECASE,
)

_EMAIL_RE = re.compile(r"[^\s@]+@[^\s@]+\.[^\s@]+")
_URL_RE = re.compile(r"https?://|www\.", re.IGNORECASE)
_PHONE_RE = re.compile(
    r"(\+?\d[\d\s().-]{6,}\d)|(\(\d{3}\)\s*\d{3}[-.\s]?\d{4})"
)
_MANY_DIGITS_RE = re.compile(r"\d.*\d.*\d")


class DocumentExtractionError(Exception):
    """Safe, client-facing extraction failure."""

    def __init__(self, message: str, status_code: int) -> None:
        super().__init__(message)
        self.message = message
        self.status_code = status_code


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    suggested_name: str
    skills: list[str]
    source_kind: str


def _extension_from_filename(filename: str | None) -> str | None:
    if not filename:
        return None
    suffix = Path(filename).suffix.lower()
    return suffix if suffix else None


def validate_uploaded_file(
    filename: str | None,
    file_bytes: bytes,
    content_type: str | None = None,
) -> str:
    """
    Validate file size and extension. Returns source kind (pdf, docx, txt, md).

    MIME type may support validation but extension is authoritative.
    """
    if len(file_bytes) == 0:
        raise DocumentExtractionError(
            "The file is empty. Choose a document with content or paste the text instead.",
            422,
        )

    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise DocumentExtractionError(
            "The file is too large. Use a file up to 5 MB or paste the text instead.",
            413,
        )

    extension = _extension_from_filename(filename)
    if extension is None or extension not in SUPPORTED_EXTENSIONS:
        raise DocumentExtractionError(
            "Unsupported file type. Upload PDF, DOCX, TXT, or MD, or paste the text instead.",
            415,
        )

    if content_type:
        lowered = content_type.lower().strip()
        mime_ok = (
            extension == ".pdf"
            and "pdf" in lowered
            or extension == ".docx"
            and (
                "wordprocessingml" in lowered
                or lowered
                in {
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "application/msword",
                }
            )
            or extension in {".txt", ".md"}
            and (lowered.startswith("text/") or lowered == "application/octet-stream")
        )
        # MIME is supporting only; inconsistent browser MIME must not block valid extensions.
        _ = mime_ok

    return EXTENSION_TO_SOURCE_KIND[extension]


def normalize_extracted_text(text: str) -> str:
    """Trim, remove nulls, collapse excessive blank lines; reject if too long."""
    cleaned = text.replace("\x00", "")
    cleaned = cleaned.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = cleaned.strip()
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

    if not cleaned:
        raise DocumentExtractionError(
            "No readable text was found. Paste the content or try a different file.",
            422,
        )

    if len(cleaned) > MAX_EXTRACTED_TEXT_LENGTH:
        raise DocumentExtractionError(
            "The extracted text is too long. Shorten the document or paste a shorter excerpt.",
            413,
        )

    return cleaned


def _decode_plain_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentExtractionError(
        "Could not read that file. Use UTF-8 text or paste the content instead.",
        422,
    )


def _extract_pdf_text(stream: BinaryIO) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentExtractionError(
            "Document extraction is temporarily unavailable. Please try again shortly.",
            500,
        ) from exc

    try:
        reader = PdfReader(stream)
    except Exception as exc:
        raise DocumentExtractionError(
            "Could not read that PDF. Try a different file or paste the text instead.",
            422,
        ) from exc

    if reader.is_encrypted:
        try:
            unlocked = reader.decrypt("")
        except Exception:
            unlocked = 0
        if unlocked == 0:
            raise DocumentExtractionError(
                "This PDF is password-protected. Remove the password or paste the text instead.",
                422,
            )

    page_count = len(reader.pages)
    if page_count > MAX_PDF_PAGES:
        raise DocumentExtractionError(
            f"This PDF has too many pages (limit {MAX_PDF_PAGES}). "
            "Use a shorter document or paste the text instead.",
            413,
        )

    parts: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text.strip():
            parts.append(page_text)

    combined = "\n".join(parts)
    if not _has_meaningful_text(combined):
        raise DocumentExtractionError(
            "This PDF does not appear to contain readable text. "
            "Upload a text-based PDF or DOCX file, or paste the text instead.",
            422,
        )
    return combined


def _extract_docx_text(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentExtractionError(
            "Document extraction is temporarily unavailable. Please try again shortly.",
            500,
        ) from exc

    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise DocumentExtractionError(
            "Could not read that document. Try a different DOCX file or paste the text instead.",
            422,
        ) from exc

    seen_paragraphs: set[str] = set()
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text and text not in seen_paragraphs:
            seen_paragraphs.add(text)
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in seen_paragraphs:
                    seen_paragraphs.add(text)
                    parts.append(text)

    combined = "\n".join(parts)
    if not combined.strip():
        raise DocumentExtractionError(
            "No readable text was found in that document. Paste the content or try another file.",
            422,
        )
    return combined


def _has_meaningful_text(text: str) -> bool:
    return bool(re.search(r"[A-Za-z0-9]", text))


def extract_text_from_bytes(
    file_bytes: bytes,
    extension: str,
) -> str:
    """Extract raw text from supported file bytes by extension."""
    if extension == ".txt" or extension == ".md":
        return _decode_plain_text(file_bytes)
    if extension == ".pdf":
        return _extract_pdf_text(BytesIO(file_bytes))
    if extension == ".docx":
        return _extract_docx_text(file_bytes)
    raise DocumentExtractionError(
        "Unsupported file type. Upload PDF, DOCX, TXT, or MD, or paste the text instead.",
        415,
    )


def _line_looks_like_name(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 60:
        return False
    lowered = stripped.lower()
    if lowered in _HEADING_IGNORE:
        return False
    if _EMAIL_RE.search(stripped) or _URL_RE.search(stripped) or _PHONE_RE.search(stripped):
        return False
    if _MANY_DIGITS_RE.search(stripped):
        return False
    if re.search(r"\d{5}", stripped):
        return False

    words = stripped.split()
    if len(words) < 2 or len(words) > 4:
        return False

    allowed = re.compile(r"^[A-Za-z][A-Za-z.'-]*$")
    if not all(allowed.match(word) for word in words):
        return False

    title_like = sum(1 for word in words if word[0].isupper() or word.isupper())
    return title_like >= max(1, len(words) - 1)


def suggest_profile_name(text: str, filename: str | None = None) -> str:
    """Deterministic editable name suggestion from early lines or filename."""
    for line in text.splitlines()[:12]:
        candidate = line.strip()
        if candidate and _line_looks_like_name(candidate):
            return candidate

    if filename:
        stem = Path(filename).stem
        cleaned = re.sub(r"[_\-]+", " ", stem)
        cleaned = re.sub(r"[^\w\s.'-]", " ", cleaned, flags=re.UNICODE)
        cleaned = _FILENAME_GENERIC_WORDS.sub(" ", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        if cleaned:
            return cleaned.title()

    return "Resume profile"


def flatten_detected_skills(skills_by_category: dict[str, list[str]]) -> list[str]:
    seen: set[str] = set()
    flattened: list[str] = []
    for skill_list in skills_by_category.values():
        for skill in skill_list:
            key = skill.lower()
            if key not in seen:
                seen.add(key)
                flattened.append(skill)
    flattened.sort(key=lambda value: value.lower())
    return flattened


def extract_taxonomy_skills(
    text: str,
    taxonomy_path: Path = _DEFAULT_TAXONOMY,
    aliases_path: Path = _DEFAULT_ALIASES,
) -> list[str]:
    """Reuse find_skills with project taxonomy and aliases."""
    taxonomy = load_json_file(taxonomy_path)
    aliases = load_json_file(aliases_path)
    found = find_skills(text, taxonomy, aliases)
    return flatten_detected_skills(found)


def extract_from_file_bytes(
    file_bytes: bytes,
    filename: str | None,
    content_type: str | None = None,
) -> ExtractionResult:
    source_kind = validate_uploaded_file(filename, file_bytes, content_type)
    extension = _extension_from_filename(filename)
    assert extension is not None

    raw_text = extract_text_from_bytes(file_bytes, extension)
    normalized = normalize_extracted_text(raw_text)
    suggested_name = suggest_profile_name(normalized, filename)
    skills = extract_taxonomy_skills(normalized)

    return ExtractionResult(
        text=normalized,
        suggested_name=suggested_name,
        skills=skills,
        source_kind=source_kind,
    )


def extract_from_pasted_text(
    text: str,
    filename: str | None = None,
) -> ExtractionResult:
    normalized = normalize_extracted_text(text)
    suggested_name = suggest_profile_name(normalized, filename)
    skills = extract_taxonomy_skills(normalized)
    return ExtractionResult(
        text=normalized,
        suggested_name=suggested_name,
        skills=skills,
        source_kind="pasted",
    )
