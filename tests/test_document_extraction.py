# Tests for api/document_extraction.py and POST /extract-document.
import os
import sys
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

repo_root = Path(__file__).resolve().parent.parent
if str(repo_root) not in sys.path:
    sys.path.insert(0, str(repo_root))

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter

import api.main as api_main_module
from api.document_extraction import (
    MAX_FILE_SIZE_BYTES,
    MAX_PDF_PAGES,
    DocumentExtractionError,
    extract_from_file_bytes,
    extract_from_pasted_text,
    extract_taxonomy_skills,
    normalize_extracted_text,
    suggest_profile_name,
    validate_uploaded_file,
)
from api.main import app

client = TestClient(app)

TEST_SHARED_SECRET = "test-analysis-api-shared-secret"
FICTIONAL_NAME = "Alex Morgan"
FICTIONAL_RESUME_BODY = (
    f"{FICTIONAL_NAME}\n"
    "Software intern experience with Python, SQL, and FastAPI.\n"
    "Built APIs and wrote technical documentation."
)
PRIVATE_MARKER = "PRIVATE_RESUME_MARKER_DO_NOT_ECHO_IN_ERRORS"


def _reload_api_app() -> TestClient:
    importlib = __import__("importlib")
    reloaded = importlib.reload(api_main_module)
    return TestClient(reloaded.app)


def _make_txt_bytes(text: str) -> bytes:
    return text.encode("utf-8")


def _make_md_bytes(text: str) -> bytes:
    return text.encode("utf-8")


def _make_docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, cell_text in enumerate(row):
                table.rows[row_index].cells[col_index].text = cell_text
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_pdf_bytes(page_count: int = 1, include_text: bool = True) -> bytes:
    if include_text and page_count == 1:
        # Minimal valid PDF with a simple text drawing operator.
        return (
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R"
            b"/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
            b"4 0 obj<</Length 55>>stream\n"
            b"BT /F1 12 Tf 72 720 Td (Python SQL FastAPI) Tj ET\n"
            b"endstream\nendobj\n"
            b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"xref\n0 6\n"
            b"0000000000 65535 f \n"
            b"0000000009 00000 n \n"
            b"0000000052 00000 n \n"
            b"0000000101 00000 n \n"
            b"0000000244 00000 n \n"
            b"0000000351 00000 n \n"
            b"trailer<</Size 6/Root 1 0 R>>\n"
            b"startxref\n408\n%%EOF\n"
        )

    writer = PdfWriter()
    for _ in range(page_count):
        writer.add_blank_page(width=612, height=792)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _assert_safe_error_text(text: str) -> None:
    for marker in [PRIVATE_MARKER, "Traceback", ".py", TEST_SHARED_SECRET]:
        assert marker not in text


def test_normalize_extracted_text_trims_and_collapses_blank_lines():
    result = normalize_extracted_text("  line one\n\n\n\nline two  ")
    assert result == "line one\n\nline two"


def test_normalize_extracted_text_rejects_empty():
    try:
        normalize_extracted_text("   ")
        assert False, "expected error"
    except DocumentExtractionError as exc:
        assert exc.status_code == 422


def test_normalize_extracted_text_rejects_too_long():
    try:
        normalize_extracted_text("a" * 100_001)
        assert False, "expected error"
    except DocumentExtractionError as exc:
        assert exc.status_code == 413


def test_txt_extraction():
    result = extract_from_file_bytes(_make_txt_bytes(FICTIONAL_RESUME_BODY), "sample.txt")
    assert result.source_kind == "txt"
    assert "Python" in result.text
    assert result.skills


def test_md_extraction():
    result = extract_from_file_bytes(_make_md_bytes(FICTIONAL_RESUME_BODY), "sample.md")
    assert result.source_kind == "md"
    assert "FastAPI" in result.text


def test_docx_paragraph_extraction():
    payload = _make_docx_bytes([FICTIONAL_NAME, "Experience with Python and SQL."])
    result = extract_from_file_bytes(payload, "profile.docx")
    assert result.source_kind == "docx"
    assert FICTIONAL_NAME in result.text


def test_docx_table_cell_extraction():
    payload = _make_docx_bytes(
        [FICTIONAL_NAME],
        table_rows=[["Skill", "Detail"], ["Python", "APIs"], ["SQL", "Queries"]],
    )
    result = extract_from_file_bytes(payload, "table-profile.docx")
    assert "Python" in result.text
    assert "Queries" in result.text


def test_pdf_text_extraction():
    payload = _make_pdf_bytes(include_text=True)
    result = extract_from_file_bytes(payload, "skills.pdf")
    assert result.source_kind == "pdf"
    assert "Python" in result.text or "SQL" in result.text


def test_scanned_or_empty_pdf_rejected():
    payload = _make_pdf_bytes(include_text=False)
    try:
        extract_from_file_bytes(payload, "blank.pdf")
        assert False, "expected error"
    except DocumentExtractionError as exc:
        assert exc.status_code == 422
        assert "readable text" in exc.message.lower()


def test_unsupported_extension_rejected():
    try:
        validate_uploaded_file("notes.rtf", b"hello")
        assert False, "expected error"
    except DocumentExtractionError as exc:
        assert exc.status_code == 415


def test_oversized_file_rejected():
    try:
        validate_uploaded_file("big.txt", b"x" * (MAX_FILE_SIZE_BYTES + 1))
        assert False, "expected error"
    except DocumentExtractionError as exc:
        assert exc.status_code == 413


def test_too_many_pdf_pages_rejected():
    payload = _make_pdf_bytes(page_count=MAX_PDF_PAGES + 1, include_text=False)
    try:
        extract_from_file_bytes(payload, "long.pdf")
        assert False, "expected error"
    except DocumentExtractionError as exc:
        assert exc.status_code == 413


def test_malformed_docx_rejected():
    try:
        extract_from_file_bytes(b"not-a-real-docx", "broken.docx")
        assert False, "expected error"
    except DocumentExtractionError as exc:
        assert exc.status_code == 422
        _assert_safe_error_text(exc.message)


def test_pasted_text_extraction():
    result = extract_from_pasted_text(FICTIONAL_RESUME_BODY)
    assert result.source_kind == "pasted"
    assert "Python" in result.skills or "SQL" in result.skills


def test_suggested_name_from_early_line():
    assert suggest_profile_name(FICTIONAL_RESUME_BODY) == FICTIONAL_NAME


def test_suggested_name_filename_fallback():
    name = suggest_profile_name("Skills\nPython", "alex_morgan_resume_final.pdf")
    assert "Alex" in name


def test_suggested_name_final_fallback():
    assert suggest_profile_name("Skills\nPython", None) == "Resume profile"


def test_deterministic_skill_extraction():
    skills = extract_taxonomy_skills("Built services with Python, SQL, and FastAPI.")
    lowered = {skill.lower() for skill in skills}
    assert "python" in lowered
    assert "sql" in lowered
    assert "fastapi" in lowered


def test_extract_document_endpoint_success():
    response = client.post(
        "/extract-document",
        data={"text": FICTIONAL_RESUME_BODY},
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["sourceKind"] == "pasted"
    assert payload["suggestedName"]
    assert isinstance(payload["skills"], list)
    assert PRIVATE_MARKER not in response.text


def test_extract_document_rejects_both_file_and_text():
    response = client.post(
        "/extract-document",
        data={"text": FICTIONAL_RESUME_BODY},
        files={"file": ("sample.txt", _make_txt_bytes("other"), "text/plain")},
    )
    assert response.status_code == 400
    _assert_safe_error_text(response.text)


def test_extract_document_rejects_neither_input():
    response = client.post("/extract-document", data={})
    assert response.status_code == 400


def test_extract_document_file_upload_txt():
    response = client.post(
        "/extract-document",
        files={"file": ("sample.txt", _make_txt_bytes(FICTIONAL_RESUME_BODY), "text/plain")},
    )
    assert response.status_code == 200
    assert response.json()["sourceKind"] == "txt"


def test_extract_document_unsupported_extension_returns_415():
    response = client.post(
        "/extract-document",
        files={"file": ("notes.rtf", b"hello", "application/rtf")},
    )
    assert response.status_code == 415
    _assert_safe_error_text(response.text)


def test_extract_document_oversized_returns_413():
    response = client.post(
        "/extract-document",
        files={"file": ("big.txt", b"x" * (MAX_FILE_SIZE_BYTES + 1), "text/plain")},
    )
    assert response.status_code == 413


def test_extract_document_unreadable_pdf_returns_422():
    response = client.post(
        "/extract-document",
        files={"file": ("blank.pdf", _make_pdf_bytes(include_text=False), "application/pdf")},
    )
    assert response.status_code == 422
    _assert_safe_error_text(response.text)


def test_extract_document_shared_secret_protection():
    original = os.environ.get("ANALYSIS_API_SHARED_SECRET")
    os.environ["ANALYSIS_API_SHARED_SECRET"] = TEST_SHARED_SECRET
    try:
        reloaded_client = _reload_api_app()
        missing = reloaded_client.post(
            "/extract-document",
            data={"text": FICTIONAL_RESUME_BODY},
        )
        assert missing.status_code == 401

        authorized = reloaded_client.post(
            "/extract-document",
            data={"text": FICTIONAL_RESUME_BODY},
            headers={"X-Analysis-Api-Key": TEST_SHARED_SECRET},
        )
        assert authorized.status_code == 200
    finally:
        if original is None:
            os.environ.pop("ANALYSIS_API_SHARED_SECRET", None)
        else:
            os.environ["ANALYSIS_API_SHARED_SECRET"] = original
        _reload_api_app()


def test_extract_document_safe_500_without_exception_details():
    with patch(
        "api.main.extract_from_pasted_text",
        side_effect=RuntimeError(PRIVATE_MARKER),
    ):
        response = client.post(
            "/extract-document",
            data={"text": FICTIONAL_RESUME_BODY},
        )
    assert response.status_code == 500
    _assert_safe_error_text(response.text)
